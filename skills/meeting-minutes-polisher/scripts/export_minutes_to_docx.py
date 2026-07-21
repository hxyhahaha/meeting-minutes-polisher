#!/usr/bin/env python3
"""
Export polished meeting minutes to a Word document.

Usage:
    python export_minutes_to_docx.py --company "公司名" --input /path/to/minutes.txt --output /path/to/output.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def build_document(company: str, content: str) -> Document:
    document = Document()

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f"{company}会议纪要")
    title_run.bold = True
    title_run.font.size = Pt(16)

    document.add_paragraph("")

    for block in content.split("\n\n"):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(block.strip())

    return document


def main() -> None:
    parser = argparse.ArgumentParser(description="Export meeting minutes to DOCX.")
    parser.add_argument("--company", required=True, help="Company name for the document title")
    parser.add_argument("--input", required=True, help="Path to the polished minutes text file")
    parser.add_argument("--output", required=True, help="Path to the output DOCX file")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)

    content = input_path.read_text(encoding="utf-8").strip()
    document = build_document(args.company, content)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    main()
